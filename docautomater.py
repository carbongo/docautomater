import os
import PySimpleGUI as sg
import pandas as pd
import docx

def read_table_file(table_file):
    if table_file.endswith('.csv'):
        return pd.read_csv(table_file)
    elif table_file.endswith('.xlsx'):
        return pd.read_excel(table_file)
    else:
        sg.popup_error('Invalid table file format. Please select a csv or xlsx file.')
        return None

def read_doc_file(doc_file):
    if doc_file.endswith('.docx'):
        return docx.Document(doc_file)
    else:
        sg.popup_error('Invalid document file format. Please select a docx file.')
        return None

def replace_keywords(doc, person_info):
    for p in doc.paragraphs:
        for key, value in person_info.items():
            if key in p.text:
                p.text = p.text.replace(key, str(value))
    return doc

def save_modified_doc(doc, doc_file, person_info, output_folder):
    new_doc_file = os.path.join(output_folder, os.path.basename(doc_file).split('.')[0] + ' ' + list(person_info.values())[0] + '.' + doc_file.split('.')[1])
    doc.save(new_doc_file)

# Define PySimpleGUI layout
layout = [
    [sg.Text('Select table file (csv or xlsx):'), sg.Input(key='table_file'), sg.FileBrowse()],
    [sg.Text('Select document file (docx):'), sg.Input(key='doc_file'), sg.FileBrowse()],
    [sg.Text('Select output folder:'), sg.Input(key='output_folder', default_text=os.getcwd()), sg.FolderBrowse()],
    [sg.Button('Generate')]
]

# Create PySimpleGUI window
window = sg.Window('Document Automater', layout)

# Event loop
while True:
    event, values = window.read()
    if event == sg.WIN_CLOSED:
        break
    if event == 'Generate':
        # Read in table file
        table_file = values['table_file']
        table_df = read_table_file(table_file)
        if table_df is None:
            continue
        
        # Read in document file
        doc_file = values['doc_file']
        doc = read_doc_file(doc_file)
        if doc is None:
            continue
        
        # Replace keywords with person's information and save modified document file
        output_folder = values['output_folder']
        for i, row in table_df.iterrows():
            person_info = dict(row)
            doc = docx.Document(doc_file)
            modified_doc = replace_keywords(doc, person_info)
            save_modified_doc(modified_doc, doc_file, person_info, output_folder)
        
        sg.popup('Files generated successfully.')
